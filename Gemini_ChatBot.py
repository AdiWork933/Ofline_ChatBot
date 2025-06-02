import os
import re
import json
import string
import inflect
import nltk
import pandas as pd
import requests
import streamlit as st
from datetime import datetime
from base64 import b64encode
from dotenv import load_dotenv
from fuzzywuzzy import fuzz
from PyPDF2 import PdfReader
import google.generativeai as genai
from docx import Document as DocxDocument
from transformers import AutoTokenizer
import torch
from sentence_transformers import util

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer, PorterStemmer

# Download required NLTK data
nltk.download('punkt')
nltk.download('wordnet')
nltk.download('stopwords')

# Load Gemini API Key
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Streamlit settings
st.set_page_config(page_title="EDU GEN", page_icon="img/edugen.png", layout="wide")

# Initialize session state variables before any usage
for key in ["chat_history", "show_feedback", "pdf_text"]:
    if key not in st.session_state:
        st.session_state[key] = [] if key == "chat_history" else False


# Custom CSS to remove red border on input focus
st.markdown("""
    <style>
    /* Universal reset for Streamlit input fields and buttons */
    .stTextInput input[type="text"],
    .stTextInput input[type="text"]:focus,
    .stTextInput input[type="text"]:hover,
    .stTextInput input[type="text"]:active,
    .stButton > button,
    .stButton > button:focus,
    .stButton > button:hover,
    .stButton > button:active {
        border: 1px solid #ccc !important;
        background-color: white !important;
        box-shadow: none !important;
        outline: none !important;
        color: black !important;
    }

    /* Prevent red border due to browser validation (light & dark themes) */
    input:invalid,
    input:focus:invalid,
    .stTextInput input:invalid,
    .stTextInput input:focus:invalid {
        border: 1px solid #ccc !important;
        box-shadow: none !important;
    }

    /* Optional: Button cursor & white hover effect */
    .stButton > button:hover {
        background-color: white !important;
        color: black !important;
    }
    </style>
""", unsafe_allow_html=True)


# Text preprocessing functions
def text_lowercase(text):
    return text.lower()

def remove_numbers(text):
    return re.sub(r'\d+', '', text)

p = inflect.engine()
def convert_number(text):
    return ' '.join([p.number_to_words(word) if word.isdigit() else word for word in text.split()])

def remove_punctuation(text):
    return text.translate(str.maketrans('', '', string.punctuation))

def remove_whitespace(text):
    return " ".join(text.split())

def remove_stopwords(text):
    stop_words = set(stopwords.words("english"))
    return [word for word in word_tokenize(text) if word not in stop_words]

stemmer = PorterStemmer()
def stem_words(text):
    return [stemmer.stem(word) for word in word_tokenize(text)]

lemmatizer = WordNetLemmatizer()
def lemma_words(text):
    return [lemmatizer.lemmatize(word) for word in word_tokenize(text)]

# Combine all cleaning
def clean_and_analyze_text(text):
    original = text
    text = text_lowercase(text)
    text = convert_number(text)
    text = remove_numbers(text)
    text = remove_punctuation(text)
    text = remove_whitespace(text)
    no_stopwords = remove_stopwords(text)
    stemmed = stem_words(text)
    lemmatized = lemma_words(text)
    return {
        "original": original,
        "lowercase": text_lowercase(original),
        "converted_numbers": convert_number(original),
        "removed_numbers": remove_numbers(original),
        "removed_punctuation": remove_punctuation(original),
        "removed_whitespace": remove_whitespace(original),
        "no_stopwords": no_stopwords,
        "stemmed": stemmed,
        "lemmatized": lemmatized
    }

# Load logos
def load_logo_base64(path):
    with open(path, "rb") as f:
        return b64encode(f.read()).decode()

edugen_logo = load_logo_base64("img/edugen.png")
user_logo = load_logo_base64("img/images.png")

def chat_bubble(role, content):
    theme = st.get_option("theme.base")
    is_dark = theme == "dark"

    alignment = "row-reverse" if role == "user" else "row"
    bubble_color = "#2f6f9f" if (role == "user" and is_dark) else \
                        "#d1eaff" if role == "user" else \
                        "#444" if is_dark else "#f1f1f1"

    text_color = "#fff" if is_dark and role == "assistant" else "#000"

    logo = user_logo if role == "user" else edugen_logo
    st.markdown(f"""
        <div style='display: flex; flex-direction: {alignment}; align-items: flex-start; margin-bottom: 1rem;'>
            <img src="data:image/png;base64,{logo}" style='width: 40px; height: 40px; margin: 0 12px; border-radius: 50%;' />
            <div style='background-color: {bubble_color}; color: {text_color}; padding: 12px 18px; border-radius: 16px; max-width: 75%; font-size: 15px; line-height: 1.5;'>{content}</div>
        </div>
    """, unsafe_allow_html=True)


def format_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    return text

def chunk_text(text, tokenizer_name="gpt2", chunk_size=500):
    tokenizer = AutoTokenizer.from_pretrained(tokenizer_name)
    sentences = nltk.sent_tokenize(text)
    chunks = []
    current_chunk = ""
    current_tokens = 0
    for sentence in sentences:
        tokens = tokenizer.encode(sentence, add_special_tokens=False)
        if len(tokens) > chunk_size:
            truncated_sentence = tokenizer.decode(tokens[:chunk_size])
            chunks.append(truncated_sentence.strip())
            continue
        if current_tokens + len(tokens) <= chunk_size:
            current_chunk += sentence + " "
            current_tokens += len(tokens)
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + " "
            current_tokens = len(tokens)
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks

def chunk_docx_text(self, file, tokenizer_name="gpt2", chunk_size=500):
    doc = DocxDocument(file)
    text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
    text = self.format_text(text)
    return self.chunk_text(text, tokenizer_name, chunk_size)

def chunk_excel_text(self, file, tokenizer_name="gpt2", chunk_size=500):
    tokenizer = AutoTokenizer.from_pretrained(tokenizer_name)
    try:
        df = pd.read_excel(file, sheet_name=None)
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return []
    all_sentences = []
    for sheet_name, sheet_df in df.items():
        try:
            rows = sheet_df.astype(str).apply(lambda x: ' | '.join(x), axis=1).tolist()
            labeled_rows = [f"Sheet: {sheet_name}\n{row}" for row in rows]
            all_sentences.extend(labeled_rows)
        except Exception as e:
            print(f"Error in sheet {sheet_name}: {e}")
            continue
    all_sentences = list(dict.fromkeys(all_sentences))
    # Instead, process in the order of appearance.
    sorted_sentences = all_sentences

    chunks, current_chunk, current_tokens = [], "", 0
    for sentence in sorted_sentences:
        tokens = tokenizer.encode(sentence, add_special_tokens=False)
        if len(tokens) > chunk_size:
            truncated = tokenizer.decode(tokens[:chunk_size])
            chunks.append(truncated.strip())
            continue
        if current_tokens + len(tokens) <= chunk_size:
            current_chunk += " " + sentence
            current_tokens += len(tokens)
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence
            current_tokens = len(tokens)
    if current_chunk:
        chunks.append(current_chunk.strip())
    return list(dict.fromkeys(chunks))

def chunk_pdf_text(self, file, tokenizer_name="gpt2", chunk_size=500):
    tokenizer = AutoTokenizer.from_pretrained(tokenizer_name)
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text

    text = re.sub(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', '', text)
    text = re.sub(r'(\+?\d{1,3}[-.\s]?|\()?\d{3}[-.\s)]*\d{3}[-.\s]?\d{4}', '', text)
    text = self.format_text(text)
    return self.chunk_text(text, tokenizer_name, chunk_size)
def custom_tokenizer(text):
    """Basic word-based tokenizer"""
    return text.split()

def chunk_text(text, chunk_size):
    """Chunks text using custom tokenizer"""
    tokens = custom_tokenizer(text)
    return [
        " ".join(tokens[i:i + chunk_size])
        for i in range(0, len(tokens), chunk_size)
    ]

def process_json_file(filepath, chunk_size=100):
    """
    Loads a JSON file, chunks each object using a custom tokenizer,
    and returns a list of chunked JSON strings or parsed JSON objects.
    """
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)
            objects = data if isinstance(data, list) else [data]
            chunk_list = []

            for obj in objects:
                json_text = json.dumps(obj, separators=(",", ":"), ensure_ascii=False)
                chunks = chunk_text(json_text, chunk_size)
                chunk_list.extend([
                    json.loads(c) if c.strip().startswith("{") and c.strip().endswith("}") else c
                    for c in chunks
                ])
            return chunk_list
    except Exception as e:
        print(f"Error processing JSON file {filepath}: {e}")
        return []

def get_all_file_text():
    folder = "uploads_data"
    all_text = ""
    tokenizer = AutoTokenizer.from_pretrained("gpt2")
    chunk_size = 500

    for filename in os.listdir(folder):
        filepath = os.path.join(folder, filename)

        if filename.endswith(".pdf"):
            with open(filepath, "rb") as f:
                reader = PdfReader(f)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
                text = re.sub(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', '', text)
                text = re.sub(r'(\+?\d{1,3}[-.\s]?|\()?\d{3}[-.\s)]*\d{3}[-.\s]?\d{4}', '', text)
                text = format_text(text)
                chunks = chunk_text(text, chunk_size)  # ✅ Fixed
                all_text += "\n".join(chunks) + "\n"

        elif filename.endswith(".docx"):
            doc = DocxDocument(filepath)
            text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
            text = format_text(text)
            chunks = chunk_text(text, chunk_size)  # ✅ Fixed
            all_text += "\n".join(chunks) + "\n"

        elif filename.endswith(".xlsx"):
            try:
                df = pd.read_excel(filepath, sheet_name=None)
                all_sentences = []
                for sheet_name, sheet_df in df.items():
                    rows = sheet_df.astype(str).apply(lambda x: ' | '.join(x), axis=1).tolist()
                    labeled_rows = [f"Sheet: {sheet_name}\n{row}" for row in rows]
                    all_sentences.extend(labeled_rows)
                sorted_sentences = list(dict.fromkeys(all_sentences))

                current_chunk, current_tokens = "", 0
                for sentence in sorted_sentences:
                    tokens = tokenizer.encode(sentence, add_special_tokens=False)
                    if len(tokens) > chunk_size:
                        truncated = tokenizer.decode(tokens[:chunk_size])
                        all_text += truncated.strip() + "\n"
                        continue
                    if current_tokens + len(tokens) <= chunk_size:
                        current_chunk += " " + sentence
                        current_tokens += len(tokens)
                    else:
                        all_text += current_chunk.strip() + "\n"
                        current_chunk = sentence
                        current_tokens = len(tokens)
                if current_chunk:
                    all_text += current_chunk.strip() + "\n"
            except Exception as e:
                print(f"Error processing Excel file {filename}: {e}")

        elif filename.endswith(".json"):
            try:
                result_chunks = process_json_file(filepath, chunk_size=50)
                for chunk in result_chunks:
                    if isinstance(chunk, dict):
                        all_text += json.dumps(chunk, ensure_ascii=False) + "\n"
                    elif isinstance(chunk, str):
                        all_text += chunk + "\n"
            except Exception as e:
                print(f"Error processing JSON file {filename}: {e}")

    return all_text

def ask_gemini(context, question):
    prompt = f"""
Please analyze the following context extracted from a PDF and answer the question as thoroughly and clearly as possible.

- Do not copy full paragraphs from the context.
- Analyze and synthesize the information before responding.
- If the question is irrelevant to the context(analyze perfect irrelevant questions or texts), respond with: "I can't understand your question."
- If the question is related to admission, criteria, or requirements, provide an answer **specific to the Maharashtra board.
- in json Ai represent All india rank(merit, condider all rank available in data) give proper answer from the context.
- If the answer of FC, Rank, Application id are not available in the given context then response with: "Sorry! Your application id is out of list."

Context:
{context}
- arrange the answer in proper format(add bullet point, header, paragraphs and so on)

Question:
{question}

Answer:
"""
    model = genai.GenerativeModel("models/gemini-2.0-flash")
    response = model.generate_content(prompt)
    return response.text.strip()


def is_small_talk_or_fuzzy_match(text):
    predefined = {
        "which candidates are allowed in cap round 2 in pharmacy": "Hello! How can I help you today?",
        "hi": "Hello! How can I help you today?",
        "hello": "Hi there! What can I do for you?",
        "hey": "Hey! Need help with something?",
        "how are you": "I'm great, thanks for asking! How can I assist?",
        "thank you": "You're welcome! Do you have more questions?",
        "thanks": "No problem! Ask away.",
        "good morning": "Good morning! How can I help?",
        "good evening": "Good evening! Ready when you are.",
        "yes": "Sure, Go ahead!",
        "no": "It's okay.",
        "yo": "Yo! What's your question?"
    }
    text = text.lower().strip()
    for q in predefined:
        score = fuzz.token_sort_ratio(q, text)
        if score > 85:
            return predefined[q]
    return None

def save_to_json(question, answer, filename="chat_history.json"):
    entry = {"question": question, "answer": answer}
    if os.path.exists(filename):
        with open(filename, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = []
    data.append(entry)
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)

for key in ["chat_history", "show_feedback", "pdf_text"]:
    if key not in st.session_state:
        st.session_state[key] = [] if key == "chat_history" else False

if not st.session_state["pdf_text"]:
    st.session_state["pdf_text"] = get_all_file_text()

# st.title("EDUGEN Assistant")

for entry in st.session_state.chat_history:
    chat_bubble("user", entry["question"])
    chat_bubble("assistant", entry["answer"])
    if entry.get("feedback"):
        st.markdown(f"📝Feedback: {entry['feedback']}")

if st.session_state.get("show_feedback", False):
    with st.expander("📝Provide Feedback for the Last Answer"):
        app_id = st.text_input("Application ID")  # This will be used as course_name
        feedback_text = st.text_input("Your feedback:")

        if st.button("Submit Feedback"):
            if not app_id.strip() or not feedback_text.strip():
                st.warning("⚠ Both Application ID and Feedback are required.")
            else:
                last_qna = st.session_state.chat_history[-1]
                last_qna["feedback"] = feedback_text
                last_qna["app_id"] = app_id

                # Save to text log
                with open("feedback_log.txt", "a", encoding="utf-8") as f:
                    f.write(
                        f"{datetime.now()}\nApp ID: {app_id}\nQ: {last_qna['question']}\nA: {last_qna['answer']}\nFeedback: {feedback_text}\n{'-'*40}\n"
                    )

                # Save to JSON log
                feedback_data = {
                    "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    "application_id": app_id,
                    "question": last_qna["question"],
                    "answer": last_qna["answer"],
                    "feedback": feedback_text
                }

                if os.path.exists("feedback_log.json"):
                    with open("feedback_log.json", "r+", encoding="utf-8") as jf:
                        try:
                            data = json.load(jf)
                        except json.JSONDecodeError:
                            data = []
                        data.append(feedback_data)
                        jf.seek(0)
                        json.dump(data, jf, indent=4)
                else:
                    with open("feedback_log.json", "w", encoding="utf-8") as jf:
                        json.dump([feedback_data], jf, indent=4)

                # Send to API (corrected payload format)
                try:
                    api_payload = {
                        "username": app_id,
                        "question": last_qna["question"],
                        "feedback": feedback_text
                    }
                    headers = {"Content-Type": "application/json"}
                    response = requests.post(
                        "https://m4qv25rg-9999.inc1.devtunnels.ms/miscellaneous/feedbackFromChatBot",
                        json=api_payload,
                        headers=headers,
                        timeout=10
                    )

                    if response.status_code == 200 and "success" in response.text.lower():
                        last_qna["api_feedback_status"] = "submitted"
                        st.success("✅ Feedback submitted successfully!")
                    else:
                        last_qna["api_feedback_status"] = f"failed ({response.status_code})"
                        st.warning(f"⚠ Feedback saved locally but failed to submit to the server. Server said: {response.text}")
                except Exception as e:
                    last_qna["api_feedback_status"] = f"failed (exception)"
                    st.warning(f"⚠ Failed to send feedback to server. Error: {e}")

                st.session_state.show_feedback = False

user_question = st.chat_input("Type your question here...")
if st.button("📝Give Feedback"):
    st.session_state.show_feedback = True

if user_question:
    with st.spinner("Generating answer..."):
        response = is_small_talk_or_fuzzy_match(user_question)
        if response:
            answer = response
        elif not st.session_state["pdf_text"]:
            answer = "⚠ Please ensure PDF, DOCX, and XLSX files are present in the uploads_data folder."
        else:
            answer = ask_gemini(st.session_state["pdf_text"], user_question)

        chat_bubble("user", user_question)
        chat_bubble("assistant", answer)

        st.session_state.chat_history.append({
            "question": user_question,
            "answer": answer,
            "feedback": None
        })

        save_to_json(user_question, answer)
        st.session_state.show_feedback = False
        
#Running port :--https://9x2g83pl-8505.inc1.devtunnels.ms/
