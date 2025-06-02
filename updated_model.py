import re 
import nltk
import os
import torch
import pandas as pd
from PyPDF2 import PdfReader
from nltk.tokenize import sent_tokenize
from docx import Document as DocxDocument
from typing import List
from transformers import (
    T5ForConditionalGeneration,
    T5Tokenizer,
    AutoTokenizer,
    AutoModelForCausalLM,
    BitsAndBytesConfig
)
from sentence_transformers import SentenceTransformer, util
from langchain.docstore.document import Document
from langchain.vectorstores import Chroma
from langchain.embeddings import HuggingFaceEmbeddings
from rapidfuzz import fuzz

nltk.download('punkt')

class ModelHandler:
    def __init__(self):
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

        self.t5_model = T5ForConditionalGeneration.from_pretrained("google/flan-t5-base").to(self.device)
        self.t5_tokenizer = T5Tokenizer.from_pretrained("google/flan-t5-base")

        self.embedding_model = SentenceTransformer("all-MiniLM-L6-v2", device=self.device)
        self.embedding_model_chroma = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

        self.chunk_tokenizer = AutoTokenizer.from_pretrained("gpt2")

        self.chroma_persist_dir = "chroma_store"
        os.makedirs(self.chroma_persist_dir, exist_ok=True)

        try:
            llama_model_id = "TinyLlama/TinyLlama-1.1B-Chat-v1.0"
            quant_config = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_use_double_quant=True,
                bnb_4bit_compute_dtype=torch.float16,
                bnb_4bit_quant_type="nf4"
            )

            self.llama_tokenizer = AutoTokenizer.from_pretrained(llama_model_id, use_fast=True)
            self.llama_model = AutoModelForCausalLM.from_pretrained(
                llama_model_id,
                quantization_config=quant_config,
                device_map="auto",
                trust_remote_code=True
            )

        except RuntimeError as e:
            print("Error loading TinyLLaMA model:", e)
            self.llama_model = None
            self.llama_tokenizer = None

    def format_text(self, text):
        text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
        text = re.sub(r'\n{2,}', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        lines = text.splitlines()
        formatted = [
            line.title() if line.isupper() and len(line.split()) < 10 else line.strip()
            for line in lines if line.strip()
        ]
        return "\n\n".join(formatted)

    def chunk_text(self, text, tokenizer_name="gpt2", chunk_size=500):
        tokenizer = AutoTokenizer.from_pretrained(tokenizer_name)
        sentences = sent_tokenize(text)
        chunks, current_chunk, current_tokens = [], "", 0
        for sentence in sentences:
            tokens = tokenizer.encode(sentence, add_special_tokens=False)
            if current_tokens + len(tokens) <= chunk_size:
                current_chunk += " " + sentence
                current_tokens += len(tokens)
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
                current_tokens = len(tokens)
        if current_chunk:
            chunks.append(current_chunk.strip())
        return list(dict.fromkeys(chunks))

    def chunk_docx_text(self, file, tokenizer_name="gpt2", chunk_size=500):
        doc = DocxDocument(file)
        text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
        text = self.format_text(text)
        return self.chunk_text(text, tokenizer_name, chunk_size)

    def chunk_excel_text(self, file, tokenizer_name="gpt2", chunk_size=500):
        tokenizer = AutoTokenizer.from_pretrained(tokenizer_name)
        try:
            df = pd.read_excel(file, sheet_name=None)
        except Exception as e:
            print(f"Error reading Excel file: {e}")
            return []
        all_sentences = []
        for sheet_name, sheet_df in df.items():
            try:
                rows = sheet_df.astype(str).apply(lambda x: ' | '.join(x), axis=1).tolist()
                labeled_rows = [f"Sheet: {sheet_name}\n{row}" for row in rows]
                all_sentences.extend(labeled_rows)
            except Exception as e:
                print(f"Error in sheet {sheet_name}: {e}")
                continue
        all_sentences = list(dict.fromkeys(all_sentences))

        embeddings = self.embedding_model.encode(all_sentences, convert_to_tensor=True)
        mean_embedding = torch.mean(embeddings, dim=0, keepdim=True)
        similarities = util.cos_sim(mean_embedding, embeddings)[0]
        ranked_indices = torch.argsort(similarities, descending=True)
        sorted_sentences = [all_sentences[i] for i in ranked_indices]

        chunks, current_chunk, current_tokens = [], "", 0
        for sentence in sorted_sentences:
            tokens = tokenizer.encode(sentence, add_special_tokens=False)
            if len(tokens) > chunk_size:
                truncated = tokenizer.decode(tokens[:chunk_size])
                chunks.append(truncated.strip())
                continue
            if current_tokens + len(tokens) <= chunk_size:
                current_chunk += " " + sentence
                current_tokens += len(tokens)
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
                current_tokens = len(tokens)
        if current_chunk:
            chunks.append(current_chunk.strip())
        return list(dict.fromkeys(chunks))

    def chunk_pdf_text(self, file, tokenizer_name="gpt2", chunk_size=500):
        tokenizer = AutoTokenizer.from_pretrained(tokenizer_name)
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text

        text = re.sub(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', '', text)
        text = re.sub(r'(\+?\d{1,3}[-.\s]?|\()?\d{3}[-.\s)]*\d{3}[-.\s]?\d{4}', '', text)
        text = self.format_text(text)
        return self.chunk_text(text, tokenizer_name, chunk_size)

    def get_vector_store_from_chunks(self, chunks: List[str]):
        if not chunks or not isinstance(chunks, list):
            raise ValueError("Chunks must be a non-empty list of strings.")
        texts = chunks
        metadatas = [{"source": f"chunk_{i}"} for i in range(len(texts))]
        vector_store = Chroma.from_texts(
            texts=texts,
            embedding=self.embedding_model_chroma,
            metadatas=metadatas,
            persist_directory=self.chroma_persist_dir
        )
        vector_store.persist()
        documents = [Document(page_content=text, metadata=metadata) for text, metadata in zip(texts, metadatas)]
        return vector_store, documents

    def enhance_with_llama(self, question, initial_answer, context):
        if self.llama_model is None or self.llama_tokenizer is None:
            return initial_answer

        should_generate = (
            not initial_answer or
            len(initial_answer.strip()) < 20 or
            len(sent_tokenize(initial_answer)) < 2
        )

        if should_generate:
            llama_prompt = (
                f"[Context]\n{context.strip()}\n\n"
                f"[Question]\n{question.strip()}\n\n"
                f"[Instruction]\n"
                f"Answer the question clearly, concisely, and point-to-point using the provided context. "
                f"Include only relevant and factual information. Use numbered or bulleted points if needed. "
                f"Avoid repetition, and do not hallucinate.\n\n[Answer]"
            )
        else:
            similarity = fuzz.token_set_ratio(question.lower(), initial_answer.lower())
            if similarity > 80:
                return initial_answer  # avoid unnecessary enhancement for high-similarity answers

            llama_prompt = (
                f"[Context]\n{context.strip()}\n\n"
                f"[Question]\n{question.strip()}\n\n"
                f"[Initial Answer]\n{initial_answer.strip()}\n\n"
                f"[Instruction]\n"
                f"Improve the initial answer based on the context. Ensure it is point-to-point, well-structured, "
                f"and includes descriptive details. Use lists or formatting when appropriate. "
                f"Avoid repetition, hallucination, or vague content.\n\n[Improved Answer]"
            )

        try:
            inputs = self.llama_tokenizer(llama_prompt, return_tensors="pt", truncation=True, max_length=1024)
            inputs = {k: v.to(self.llama_model.device) for k, v in inputs.items()}
            outputs = self.llama_model.generate(
                **inputs,
                max_new_tokens=600,
                do_sample=False,
                temperature=0.7,
                repetition_penalty=1.2
            )
            decoded = self.llama_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()

            # Extract relevant content
            if "[Answer]" in decoded:
                improved_answer = decoded.split("[Answer]", 1)[-1].strip()
            elif "[Improved Answer]" in decoded:
                improved_answer = decoded.split("[Improved Answer]", 1)[-1].strip()
            else:
                improved_answer = decoded.strip()

            # Remove duplicate lines
            lines = improved_answer.splitlines()
            seen = set()
            cleaned_lines = []
            for line in lines:
                norm = line.strip().lower()
                if norm and norm not in seen:
                    cleaned_lines.append(line.strip())
                    seen.add(norm)
            return "\n".join(cleaned_lines)

        except RuntimeError as e:
            print("Error during TinyLLaMA enhancement:", e)
            return initial_answer
       
    def generate_answer(self, question, context):
        question = question.strip()
        context = context[:1800]

        is_true_false = any(question.lower().startswith(prefix) for prefix in [
            "is", "are", "was", "were", "do", "does", "did", "can", "could", "should", "would",
            "will", "has", "have", "had"
        ])

        if is_true_false:
            prompt = (
                f"You are a fact-checking assistant. Based on the context below, answer ONLY with 'True' or 'False'. "
                f"Do not explain your answer.\n\nContext:\n{context}\n\n"
                f"Question: {question}\n\nAnswer:"
            )
            inputs = self.t5_tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024).to(self.device)
            output_ids = self.t5_model.generate(**inputs, max_new_tokens=10)
            answer = self.t5_tokenizer.decode(output_ids[0], skip_special_tokens=True).strip().lower()
            raw_answer = "True" if "true" in answer else "False" if "false" in answer else "Unable to determine."
        else:
            prompt = (
                f"You are a highly knowledgeable academic assistant. Based on the context below, write a clear and accurate answer.\n"
                f"Structure your answer in a point-to-point format using numbered or bullet points. "
                f"Avoid repetition and include only relevant, factual information.\n\n"
                f"Context:\n{context}\n\nQuestion: {question}\n\nAnswer:"
            )
            inputs = self.t5_tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024).to(self.device)
            outputs = self.t5_model.generate(**inputs, max_new_tokens=800)
            answer = self.t5_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()
            if len(sent_tokenize(answer)) < 4:
                prompt += "\n\nPlease elaborate and be more comprehensive."
                inputs = self.t5_tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024).to(self.device)
                outputs = self.t5_model.generate(**inputs, max_new_tokens=600)
                answer = self.t5_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()
            raw_answer = answer

        final_answer = self.enhance_with_llama(question, raw_answer, context)
        return final_answer

    def handle_small_talk(self, query):
        greetings = ["hi", "hello", "hey", "good morning", "good afternoon", "good evening"]
        thanks = ["thanks", "thank you", "thanks a lot", "cheers"]
        bye = ["bye", "goodbye", "see you", "later"]
        feed = ["why you want feedback","Requirement of feedback"]

        q = query.lower().strip()
        if q  in feed:
            return "I want to improve myself in order to enhance the quality of my answers."
        if q in greetings:
            return "Hello! I'm ready to help you with your documents."
        if q in thanks:
            return "You're welcome! Feel free to ask me more."
        if q in bye:
            return "Goodbye! Come back anytime."
        return None
