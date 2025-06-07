import re
import nltk
import os
import torch
import pandas as pd
from PyPDF2 import PdfReader
from nltk.tokenize import sent_tokenize
from docx import Document as DocxDocument
from typing import List
from transformers import (
    T5ForConditionalGeneration,
    T5Tokenizer,
    AutoTokenizer,
    AutoModelForCausalLM,
    BitsAndBytesConfig
)
from sentence_transformers import SentenceTransformer, util
from langchain.docstore.document import Document
from langchain.vectorstores import Chroma
from langchain.embeddings import HuggingFaceEmbeddings
from rapidfuzz import fuzz

nltk.download('punkt')

class ModelHandler:
    def __init__(self):
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

        self.t5_model = T5ForConditionalGeneration.from_pretrained("google/flan-t5-base").to(self.device)
        self.t5_tokenizer = T5Tokenizer.from_pretrained("google/flan-t5-base")

        self.embedding_model = SentenceTransformer("all-MiniLM-L6-v2", device=self.device)
        self.embedding_model_chroma = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

        self.chunk_tokenizer = AutoTokenizer.from_pretrained("gpt2")

        self.chroma_persist_dir = "chroma_store"
        os.makedirs(self.chroma_persist_dir, exist_ok=True)

        try:
            llama_model_id = "TinyLlama/TinyLlama-1.1B-Chat-v1.0"
            quant_config = BitsAndBytesConfig(
                load_in_4bit=True,
                bnb_4bit_use_double_quant=True,
                bnb_4bit_compute_dtype=torch.float16,
                bnb_4bit_quant_type="nf4"
            )

            self.llama_tokenizer = AutoTokenizer.from_pretrained(llama_model_id, use_fast=True)
            self.llama_model = AutoModelForCausalLM.from_pretrained(
                llama_model_id,
                quantization_config=quant_config,
                device_map="auto",
                trust_remote_code=True
            )

        except RuntimeError as e:
            print("Error loading TinyLLaMA model:", e)
            self.llama_model = None
            self.llama_tokenizer = None

    def format_text(self, text):
        text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)
        text = re.sub(r'\n{2,}', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        lines = text.splitlines()
        formatted = [
            line.title() if line.isupper() and len(line.split()) < 10 else line.strip()
            for line in lines if line.strip()
        ]
        return "\n\n".join(formatted)

    def chunk_text(self, text, tokenizer_name="gpt2", chunk_size=500):
        tokenizer = AutoTokenizer.from_pretrained(tokenizer_name)
        sentences = sent_tokenize(text)
        chunks, current_chunk, current_tokens = [], "", 0
        for sentence in sentences:
            tokens = tokenizer.encode(sentence, add_special_tokens=False)
            if current_tokens + len(tokens) <= chunk_size:
                current_chunk += " " + sentence
                current_tokens += len(tokens)
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
                current_tokens = len(tokens)
        if current_chunk:
            chunks.append(current_chunk.strip())
        return list(dict.fromkeys(chunks))

    def chunk_docx_text(self, file, tokenizer_name="gpt2", chunk_size=500):
        doc = DocxDocument(file)
        text = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
        text = self.format_text(text)
        return self.chunk_text(text, tokenizer_name, chunk_size)

    def chunk_excel_text(self, file, tokenizer_name="gpt2", chunk_size=500):
        tokenizer = AutoTokenizer.from_pretrained(tokenizer_name)
        try:
            df = pd.read_excel(file, sheet_name=None)
        except Exception as e:
            print(f"Error reading Excel file: {e}")
            return []
        all_sentences = []
        for sheet_name, sheet_df in df.items():
            try:
                rows = sheet_df.astype(str).apply(lambda x: ' | '.join(x), axis=1).tolist()
                labeled_rows = [f"Sheet: {sheet_name}\n{row}" for row in rows]
                all_sentences.extend(labeled_rows)
            except Exception as e:
                print(f"Error in sheet {sheet_name}: {e}")
                continue
        all_sentences = list(dict.fromkeys(all_sentences))

        embeddings = self.embedding_model.encode(all_sentences, convert_to_tensor=True)
        mean_embedding = torch.mean(embeddings, dim=0, keepdim=True)
        similarities = util.cos_sim(mean_embedding, embeddings)[0]
        ranked_indices = torch.argsort(similarities, descending=True)
        sorted_sentences = [all_sentences[i] for i in ranked_indices]

        chunks, current_chunk, current_tokens = [], "", 0
        for sentence in sorted_sentences:
            tokens = tokenizer.encode(sentence, add_special_tokens=False)
            if len(tokens) > chunk_size:
                truncated = tokenizer.decode(tokens[:chunk_size])
                chunks.append(truncated.strip())
                continue
            if current_tokens + len(tokens) <= chunk_size:
                current_chunk += " " + sentence
                current_tokens += len(tokens)
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
                current_tokens = len(tokens)
        if current_chunk:
            chunks.append(current_chunk.strip())
        return list(dict.fromkeys(chunks))

    def chunk_pdf_text(self, file, tokenizer_name="gpt2", chunk_size=500):
        tokenizer = AutoTokenizer.from_pretrained(tokenizer_name)
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text

        text = re.sub(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', '', text)
        text = re.sub(r'(\+?\d{1,3}[-.\s]?|\()?\d{3}[-.\s)]*\d{3}[-.\s]?\d{4}', '', text)
        text = self.format_text(text)
        return self.chunk_text(text, tokenizer_name, chunk_size)

    def get_vector_store_from_chunks(self, chunks: List[str]):
        if not chunks or not isinstance(chunks, list):
            raise ValueError("Chunks must be a non-empty list of strings.")
        texts = chunks
        metadatas = [{"source": f"chunk_{i}"} for i in range(len(texts))]
        vector_store = Chroma.from_texts(
            texts=texts,
            embedding=self.embedding_model_chroma,
            metadatas=metadatas,
            persist_directory=self.chroma_persist_dir
        )
        vector_store.persist()
        documents = [Document(page_content=text, metadata=metadata) for text, metadata in zip(texts, metadatas)]
        return vector_store, documents
    def enhance_with_llama(self, question, initial_answer, context):
        if self.llama_model is None or self.llama_tokenizer is None:
            return initial_answer.strip()

        question = question.strip()
        context = context.strip()
        initial_answer = initial_answer.strip()

        should_generate = (
            not initial_answer or
            len(initial_answer) < 20 or
            len(sent_tokenize(initial_answer)) < 2
        )

        if should_generate:
            llama_prompt = (
                f"[Context]\n{context}\n\n"
                f"[Question]\n{question}\n\n"
                f"[Instruction]\n"
                f"Answer naturally and clearly. If the question is random or gibberish, respond with witty humor. "
                f"Format like: (smiling) 'That's a funny one!' or (laughs) 'Nice try!'. Otherwise, give a helpful answer.\n\n"
                f"[Answer]"
            )
        else:
            similarity = fuzz.token_set_ratio(question.lower(), initial_answer.lower())
            if similarity > 80:
                return initial_answer

            llama_prompt = (
                f"[Context]\n{context}\n\n"
                f"[Question]\n{question}\n\n"
                f"[Initial Answer]\n{initial_answer}\n\n"
                f"[Instruction]\n"
                f"Refine the answer. If it's nonsense, reply sarcastically but politely. "
                f"If serious, explain concisely and add a playful tone.\n\n"
                f"[Improved Answer]"
            )

        try:
            inputs = self.llama_tokenizer(llama_prompt, return_tensors="pt", truncation=True, max_length=1024)
            inputs = {k: v.to(self.llama_model.device) for k, v in inputs.items()}
            outputs = self.llama_model.generate(
                **inputs,
                max_new_tokens=600,
                do_sample=False,
                temperature=0.7,
                repetition_penalty=1.2
            )
            decoded = self.llama_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()

            if "[Answer]" in decoded:
                improved_answer = decoded.split("[Answer]", 1)[-1].strip()
            elif "[Improved Answer]" in decoded:
                improved_answer = decoded.split("[Improved Answer]", 1)[-1].strip()
            else:
                improved_answer = decoded.strip()

            # Clean duplicates, empty lines
            lines = improved_answer.splitlines()
            seen = set()
            cleaned_lines = []
            for line in lines:
                norm = line.strip().lower()
                if norm and norm not in seen:
                    cleaned_lines.append(line.strip())
                    seen.add(norm)

            return "\n".join(cleaned_lines)

        except RuntimeError as e:
            print("Error during LLaMA enhancement:", e)
            return initial_answer


    def generate_answer(self, question, context=""):
        question = question.strip()
        context = context.strip()[:1800]

        # Filter out gibberish
        if not any(c.isalpha() for c in question):
            return "(laughs) That doesn't even qualify as a question! Did your pet bird try to code?"

        # True/False check
        is_true_false = any(question.lower().startswith(prefix) for prefix in [
            "is", "are", "was", "were", "do", "does", "did", "can", "could", "should", "would",
            "will", "has", "have", "had"
        ])

        try:
            if is_true_false:
                prompt = (
                    f"You are a sarcastic truth detector. Only respond with 'True' or 'False'. "
                    f"If the question is silly or nonsense, make a funny remark instead.\n\n"
                    f"Context:\n{context}\n\n"
                    f"Question: {question}\n\nAnswer:"
                )
                inputs = self.t5_tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024).to(self.device)
                output_ids = self.t5_model.generate(**inputs, max_new_tokens=10)
                answer = self.t5_tokenizer.decode(output_ids[0], skip_special_tokens=True).strip().lower()

                if "true" in answer:
                    raw_answer = "True"
                elif "false" in answer:
                    raw_answer = "False"
                else:
                    raw_answer = "(smirks) I honestly can't tell. You might be messing with me."
            else:
                prompt = (
                    f"You are a witty but intelligent assistant. Based on the context, answer clearly and accurately. "
                    f"Use numbered or bullet points, keep it readable and slightly humorous.\n\n"
                    f"Context:\n{context}\n\nQuestion: {question}\n\nAnswer:"
                )
                inputs = self.t5_tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024).to(self.device)
                outputs = self.t5_model.generate(**inputs, max_new_tokens=800)
                answer = self.t5_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()

                # Add extra humor/elaboration if short
                if len(sent_tokenize(answer)) < 4:
                    prompt += "\n\nPlease elaborate more with style and subtle humor."
                    inputs = self.t5_tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024).to(self.device)
                    outputs = self.t5_model.generate(**inputs, max_new_tokens=600)
                    answer = self.t5_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()

                raw_answer = answer

        except Exception as e:
            print("Error in T5 generation:", e)
            return "(sigh) Something went wrong while trying to sound smart."

        # Enhance final output
        final_answer = self.enhance_with_llama(question, raw_answer, context)
        return final_answer

    def handle_small_talk(self, query):
        greetings = ["hi", "hello", "hey", "good morning", "good afternoon", "good evening"]
        thanks = ["thanks", "thank you", "thanks a lot", "cheers","hi brother"]
        bye = ["bye", "goodbye", "see you", "later"]
        feed = ["why you want feedback", "requirement of feedback"]

        q = query.lower().strip()
        if q in feed:
            return "I want feedback so I can evolve into a wiser, funnier digital sage. ✨"
        if q in greetings:
            return "Hey there! I'm your AI with jokes and facts. Let's get nerdy!"
        if q in thanks:
            return "You're welcome! I'm fueled by data and dad jokes."
        if q in bye:
            return "Farewell, human! Until next download."
        return None
